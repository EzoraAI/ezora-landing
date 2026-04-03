#!/usr/bin/env python3
"""
Create EzoraAI Privacy Policy Word Document
Using python-docx library
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Create output directory if it doesn't exist
output_dir = "/sessions/keen-epic-bardeen/mnt/EzoraAI-Ops/EzoraAI/Legal"
os.makedirs(output_dir, exist_ok=True)

# Create document
doc = Document()

# Set default font to Arial
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

# Set up margins (1 inch on all sides)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "EzoraAI — Privacy Policy"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para_format = header_para.paragraph_format
    header_para_format.space_after = Pt(6)

    # Format header text
    for run in header_para.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    # Add footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page [PAGE]  |  Effective Date: " + datetime.now().strftime('%B %d, %Y')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para_format = footer_para.paragraph_format
    footer_para_format.space_before = Pt(6)

    for run in footer_para.runs:
        run.font.size = Pt(10)
        run.font.name = 'Arial'

# Effective date
today = datetime.now().strftime('%B %d, %Y')

# Add DRAFT notice
draft_para = doc.add_paragraph()
draft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
draft_run = draft_para.add_run("DRAFT — FOR LEGAL REVIEW")
draft_run.bold = True
draft_run.font.size = Pt(14)
draft_run.font.color.rgb = RGBColor(208, 0, 0)
draft_para.paragraph_format.space_after = Pt(16)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Privacy Policy")
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.name = 'Arial'
title.paragraph_format.space_after = Pt(6)

# Date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run("Last Updated: " + today)
date_run.font.size = Pt(12)
date_para.paragraph_format.space_after = Pt(16)

# Helper function to add numbered section heading
def add_section_heading(doc, number, title_text):
    heading = doc.add_paragraph()
    run = heading.add_run(str(number) + ". " + title_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(8)
    return heading

# Helper function to add subsection
def add_subsection(doc, number_str, title_text):
    heading = doc.add_paragraph()
    run = heading.add_run(number_str + " " + title_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    return heading

# Helper function to add bullet point
def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.space_after = Pt(4)
    return para

# Helper function to add regular paragraph
def add_text(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.5
    return para

# ============ SECTION 1: INTRODUCTION ============
add_section_heading(doc, 1, "Introduction")

add_text(doc, 'EzoraAI (we, us, our, or Company) operates the website and platform located at https://ezora.ai (the Platform) and related services. We are committed to protecting your privacy and ensuring you have a positive experience on our Platform.')

add_text(doc, "This Privacy Policy (Policy) explains:")

add_bullet(doc, "What information we collect")
add_bullet(doc, "How we use and process that information")
add_bullet(doc, "How we share your information with third parties")
add_bullet(doc, "Your rights regarding your data")
add_bullet(doc, "How we protect your information")

add_text(doc, "By accessing or using EzoraAI, you acknowledge that you have read, understood, and agree to be bound by this Privacy Policy. If you do not agree with our practices, please do not use the Platform.")

# ============ SECTION 2: INFORMATION WE COLLECT ============
add_section_heading(doc, 2, "Information We Collect")

add_text(doc, "We collect information in various ways, including information you provide directly and information collected automatically through your use of the Platform.")

add_subsection(doc, "2.1", "Account and Registration Information")
add_text(doc, "When you create an account, we collect:")
add_bullet(doc, "Full name")
add_bullet(doc, "Email address")
add_bullet(doc, "Account type/role selection (learner, business, or expert)")
add_bullet(doc, "Password (hashed and encrypted, never stored in plain text)")

add_subsection(doc, "2.2", "Profile and Professional Information")
add_text(doc, "For expert and business users, we collect additional information:")
add_bullet(doc, "Professional biography and background")
add_bullet(doc, "Skills and expertise areas")
add_bullet(doc, "Portfolio items, case studies, and project examples")
add_bullet(doc, "Profile photo or avatar")
add_bullet(doc, "Certifications, credentials, and qualifications")
add_bullet(doc, "Hourly rates or pricing information")

add_subsection(doc, "2.3", "Payment Information")
add_text(doc, "We partner with Stripe for secure payment processing. When you make a subscription payment or complete a transaction, we collect:")
add_bullet(doc, "Billing name and address")
add_bullet(doc, "Payment method (credit card, bank account, etc.)")
add_bullet(doc, "Transaction history")
add_bullet(doc, "Subscription plan details")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run1 = para.add_run("We do NOT store complete credit card numbers. ")
run1.font.name = 'Arial'
run2 = para.add_run("Payment card data is encrypted by Stripe and compliant with PCI-DSS standards.")
run2.italic = True
run2.font.name = 'Arial'

add_subsection(doc, "2.4", "Usage and Activity Data")
add_text(doc, "We automatically collect information about your interactions with the Platform:")
add_bullet(doc, "Session history and duration")
add_bullet(doc, "Pages and features accessed")
add_bullet(doc, "Search queries")
add_bullet(doc, "Clicks and interactions")
add_bullet(doc, "Coaching sessions or project engagements accessed")

add_subsection(doc, "2.5", "Device and Technical Data")
add_text(doc, "We collect technical information about your device and connection:")
add_bullet(doc, "IP address")
add_bullet(doc, "Browser type and version")
add_bullet(doc, "Operating system")
add_bullet(doc, "Device type (mobile, desktop, tablet)")
add_bullet(doc, "Approximate geographic location (city-level, derived from IP)")

add_subsection(doc, "2.6", "Communications")
add_text(doc, "We collect and retain:")
add_bullet(doc, "Support inquiries and tickets")
add_bullet(doc, "In-session chat logs between users and experts")
add_bullet(doc, "Messages, emails, and feedback")
add_bullet(doc, "Complaints or dispute records")

# ============ SECTION 3: HOW WE USE YOUR INFORMATION ============
add_section_heading(doc, 3, "How We Use Your Information")

add_text(doc, "We use your information for the following purposes:")

add_subsection(doc, "3.1", "Service Delivery")
add_bullet(doc, "Create and manage your account")
add_bullet(doc, "Enable matchmaking between learners, businesses, and experts")
add_bullet(doc, "Facilitate coaching sessions, projects, and transactions")
add_bullet(doc, "Provide customer support and technical assistance")

add_subsection(doc, "3.2", "Payments and Billing")
add_bullet(doc, "Process subscription payments and transaction fees")
add_bullet(doc, "Issue invoices and receipts")
add_bullet(doc, "Send payment reminders and billing notifications")

add_subsection(doc, "3.3", "Platform Analytics and Improvement")
add_bullet(doc, "Analyze usage patterns and platform performance")
add_bullet(doc, "Improve features, user experience, and functionality")
add_bullet(doc, "Conduct research and develop new services")

add_subsection(doc, "3.4", "Marketing Communications")
add_text(doc, "We may use your email address to send:")
add_bullet(doc, "Platform updates and service announcements")
add_bullet(doc, "Promotional offers and feature highlights")
add_bullet(doc, "Educational content and industry news")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run = para.add_run("You may opt out of marketing communications at any time by clicking Unsubscribe in any email or updating your notification preferences.")
run.italic = True
run.font.name = 'Arial'

add_subsection(doc, "3.5", "Fraud Prevention and Security")
add_bullet(doc, "Detect and prevent fraudulent transactions")
add_bullet(doc, "Investigate unauthorized access or abuse")
add_bullet(doc, "Enforce our Terms of Service and other agreements")

add_subsection(doc, "3.6", "Blockchain Credential Recording")
add_text(doc, "EzoraAI uses blockchain technology to create immutable, publicly verifiable records of expert credentials and verified achievements. See Section 5 (Blockchain Data) for details on what information is recorded on-chain.")

# ============ SECTION 4: HOW WE SHARE YOUR INFORMATION ============
add_section_heading(doc, 4, "How We Share Your Information")

add_text(doc, "We may share your information with third parties in the following circumstances:")

add_subsection(doc, "4.1", "Service Providers")
add_text(doc, "We share necessary information with third-party service providers who assist us in operating the Platform:")

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Stripe: ")
run.bold = True
para.add_run("Payment processing, subscription management, and payout distribution. Stripe maintains PCI-DSS compliance.")
para.paragraph_format.space_after = Pt(4)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Analytics Providers: ")
run.bold = True
para.add_run("Usage analytics and performance monitoring (e.g., Google Analytics). We may anonymize or aggregate data.")
para.paragraph_format.space_after = Pt(4)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Cloud Infrastructure: ")
run.bold = True
para.add_run("Hosting, storage, and backup services.")
para.paragraph_format.space_after = Pt(4)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Communication Platforms: ")
run.bold = True
para.add_run("Email delivery, support chat (Circle.io for community), and messaging systems.")
para.paragraph_format.space_after = Pt(8)

add_subsection(doc, "4.2", "Between Platform Users")
add_text(doc, "To facilitate marketplace transactions, we share relevant information between users:")
add_bullet(doc, "Learners and Businesses see expert profiles (name, bio, skills, ratings, verified credentials)")
add_bullet(doc, "Experts see limited learner/business profiles necessary for session context")
add_bullet(doc, "In-session communications are shared only with active session participants")
add_bullet(doc, "Transaction and engagement history may be visible to matched parties")

add_subsection(doc, "4.3", "Legal and Compliance")
add_text(doc, "We may disclose information when required by law or to:")
add_bullet(doc, "Comply with legal processes, court orders, or government requests")
add_bullet(doc, "Protect EzoraAI against legal liability")
add_bullet(doc, "Enforce our Terms of Service or other agreements")
add_bullet(doc, "Protect the safety, security, or rights of our users")

add_subsection(doc, "4.4", "Blockchain (Public Records)")
add_text(doc, "Expert credentials and verified achievements are recorded on a public blockchain. See Section 5 for full details on what data is and is not publicly visible.")

add_subsection(doc, "4.5", "Business Transfers")
add_text(doc, "If EzoraAI is involved in a merger, acquisition, bankruptcy, or asset sale, your information may be transferred as part of that transaction. We will provide notice before your data becomes subject to a different privacy policy.")

# ============ SECTION 5: BLOCKCHAIN DATA ============
add_section_heading(doc, 5, "Blockchain Data")

add_text(doc, "A core differentiator of EzoraAI is the use of blockchain technology to create immutable, on-chain records of verified expert credentials. This section explains what data is recorded and what remains private.")

add_subsection(doc, "5.1", "What Gets Recorded On-Chain (Public)")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run1 = para.add_run("The following information ")
run1.font.name = 'Arial'
run2 = para.add_run("may")
run2.italic = True
run2.font.name = 'Arial'
run3 = para.add_run(" be recorded on a public blockchain:")
run3.font.name = 'Arial'

add_bullet(doc, "Expert wallet address or pseudonymous identifier")
add_bullet(doc, "Name or professional pseudonym")
add_bullet(doc, "Verified certifications and credentials")
add_bullet(doc, "Number of completed sessions/projects (anonymized count)")
add_bullet(doc, "Aggregate ratings or performance metrics")
add_bullet(doc, "Proof-of-work metadata (transaction hash, timestamp)")

add_subsection(doc, "5.2", "What Remains Private (Not On-Chain)")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run1 = para.add_run("The following sensitive information is ")
run1.font.name = 'Arial'
run2 = para.add_run("NOT")
run2.bold = True
run2.font.name = 'Arial'
run3 = para.add_run(" recorded on the public blockchain:")
run3.font.name = 'Arial'

add_bullet(doc, "Email addresses and contact information")
add_bullet(doc, "Payment information or banking details")
add_bullet(doc, "Session chat logs or conversations")
add_bullet(doc, "Portfolio items or detailed project descriptions")
add_bullet(doc, "IP addresses or device identifiers")
add_bullet(doc, "Geographic location or personal identifiers")

add_subsection(doc, "5.3", "Immutability and Permanence")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run1 = para.add_run("Once credential records are written to the blockchain, they are ")
run1.font.name = 'Arial'
run2 = para.add_run("immutable and cannot be deleted or modified.")
run2.bold = True
run2.font.name = 'Arial'

add_text(doc, "This is by design—it ensures the integrity and trust of verified expert credentials. If you have concerns about what information is being recorded on-chain, please contact us at privacy@ezora.ai before it is recorded.")

# ============ SECTION 6: DATA RETENTION ============
add_section_heading(doc, 6, "Data Retention")

add_text(doc, "We retain your information for as long as necessary to provide services, comply with legal obligations, resolve disputes, and enforce agreements:")

add_subsection(doc, "6.1", "Active Account Data")
add_text(doc, "Account, profile, and payment information is retained while your account is active and for 7 years after termination (to satisfy tax and financial audit requirements).")

add_subsection(doc, "6.2", "Session and Communication Logs")
add_text(doc, "In-session chat logs and communications are retained for 3 years after session completion for dispute resolution and quality assurance.")

add_subsection(doc, "6.3", "Technical and Usage Data")
add_text(doc, "Log files, analytics data, and technical information are typically retained for 12 months.")

add_subsection(doc, "6.4", "Account Deletion")
add_text(doc, "If you delete your account, we will remove or anonymize your personal data from our active systems within 30 days, except where required to maintain records for legal, tax, or contractual purposes. Blockchain records cannot be deleted and will remain immutable.")

# ============ SECTION 7: YOUR RIGHTS ============
add_section_heading(doc, 7, "Your Rights")

add_text(doc, "Depending on your location, you may have certain rights regarding your personal information:")

add_subsection(doc, "7.1", "Right to Access")
add_text(doc, "You have the right to request a copy of the personal information we hold about you. To submit a request, email privacy@ezora.ai with Data Access Request in the subject line.")

add_subsection(doc, "7.2", "Right to Correction")
add_text(doc, "You may correct inaccurate or incomplete information by logging into your account or contacting us.")

add_subsection(doc, "7.3", "Right to Deletion (Right to be Forgotten)")
add_text(doc, "You may request deletion of your personal data, subject to legal retention requirements. We will process requests within 30 days. However, deletion requests do not extend to blockchain records, which are immutable.")

add_subsection(doc, "7.4", "Right to Data Portability")
add_text(doc, "You may request a structured, machine-readable copy of your personal data. We will provide this in a common format (e.g., CSV or JSON) within 30 days.")

add_subsection(doc, "7.5", "Right to Opt-Out of Marketing")
add_text(doc, "You may unsubscribe from promotional emails at any time by clicking Unsubscribe or updating your preferences in your account settings.")

add_subsection(doc, "7.6", "GDPR Rights (EU/EEA Users)")
add_text(doc, "If you are in the European Union, you also have the right to lodge a complaint with your local Data Protection Authority.")

add_subsection(doc, "7.7", "CCPA Rights (California Residents)")
add_text(doc, "California residents have additional rights under the California Consumer Privacy Act (CCPA), including the right to know, delete, and opt-out of the sale of personal information (though we do not sell personal information).")

# ============ SECTION 8: COOKIES AND TRACKING ============
add_section_heading(doc, 8, "Cookies and Tracking")

add_text(doc, "We use cookies and similar tracking technologies to enhance your experience:")

add_subsection(doc, "8.1", "Types of Cookies")

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Essential Cookies: ")
run.bold = True
para.add_run("Required for login, session management, and security.")
para.paragraph_format.space_after = Pt(4)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Performance Cookies: ")
run.bold = True
para.add_run("Collect information about how you use the Platform (via Google Analytics).")
para.paragraph_format.space_after = Pt(4)

para = doc.add_paragraph(style='List Bullet')
run = para.add_run("Preference Cookies: ")
run.bold = True
para.add_run("Remember your settings and preferences.")
para.paragraph_format.space_after = Pt(8)

add_subsection(doc, "8.2", "Managing Cookies")
add_text(doc, "Most web browsers allow you to refuse cookies or alert you when cookies are being sent. You can also opt out of Google Analytics tracking through the Google Analytics opt-out browser extension.")

# ============ SECTION 9: THIRD-PARTY SERVICES ============
add_section_heading(doc, 9, "Third-Party Services")

add_text(doc, "Our Platform integrates with the following third-party services. Their privacy practices are governed by their own privacy policies:")

add_subsection(doc, "9.1", "Stripe")
add_text(doc, "Payment processing and subscription management. See Stripe's Privacy Policy at https://stripe.com/privacy.")

add_subsection(doc, "9.2", "Google Analytics")
add_text(doc, "Platform usage analytics and performance metrics. See Google's Privacy Policy at https://policies.google.com/privacy.")

add_subsection(doc, "9.3", "Circle.io")
add_text(doc, "Community platform and messaging. See Circle.io's Privacy Policy at https://circle.so/privacy.")

# ============ SECTION 10: CHILDREN'S PRIVACY ============
add_section_heading(doc, 10, "Children's Privacy")

add_text(doc, "EzoraAI is not intended for individuals under 18 years of age. We do not knowingly collect personal information from children under 18. If we become aware that we have collected information from a minor, we will delete it promptly. If you believe we have collected information from a child, please contact us at privacy@ezora.ai.")

# ============ SECTION 11: SECURITY MEASURES ============
add_section_heading(doc, 11, "Security Measures")

add_text(doc, "We implement industry-standard technical and organizational security measures to protect your data:")

add_bullet(doc, "End-to-end encryption for payment data")
add_bullet(doc, "TLS/SSL encryption for data in transit")
add_bullet(doc, "Encrypted storage for sensitive data at rest")
add_bullet(doc, "Secure password hashing and salting")
add_bullet(doc, "Regular security audits and penetration testing")
add_bullet(doc, "Stripe's PCI-DSS Level 1 compliance for payment data")
add_bullet(doc, "Two-factor authentication (2FA) available for account protection")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
run = para.add_run("No method of internet transmission or storage is 100% secure. While we implement strong security measures, we cannot guarantee absolute security of your data.")
run.italic = True
run.font.name = 'Arial'

# ============ SECTION 12: INTERNATIONAL DATA TRANSFERS ============
add_section_heading(doc, 12, "International Data Transfers")

add_text(doc, "EzoraAI is based in the United States. If you are accessing our Platform from outside the US, please be aware that:")

add_bullet(doc, "Your information will be transferred to and processed in the United States")
add_bullet(doc, "US data protection laws may differ from your home country")
add_bullet(doc, "By using EzoraAI, you consent to the transfer of your information to the US")
add_bullet(doc, "For EU/EEA users, we rely on Standard Contractual Clauses (SCCs) to facilitate lawful data transfers.")

# ============ SECTION 13: CHANGES TO THIS POLICY ============
add_section_heading(doc, 13, "Changes to This Policy")

add_text(doc, "We may update this Privacy Policy from time to time to reflect changes in our practices, technology, or legal requirements. Material changes will be communicated to you via email or a prominent notice on the Platform at least 30 days before taking effect.")

add_text(doc, "Your continued use of the Platform after such notice constitutes your acceptance of the updated Privacy Policy.")

# ============ SECTION 14: CONTACT INFORMATION ============
add_section_heading(doc, 14, "Contact Information")

add_text(doc, "If you have questions, concerns, or wish to exercise any of your rights under this Privacy Policy, please contact us:")

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(6)
run1 = para.add_run("Email: ")
run1.bold = True
run1.font.name = 'Arial'
run2 = para.add_run("privacy@ezora.ai")
run2.font.name = 'Arial'

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(6)
run1 = para.add_run("Address: ")
run1.bold = True
run1.font.name = 'Arial'
run2 = para.add_run("EzoraAI, Inc.")
run2.font.name = 'Arial'

add_text(doc, "We will respond to requests within 30 business days.")

# ============ FOOTER ============
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(16)
para.paragraph_format.space_after = Pt(8)
run = para.add_run("———")
run.italic = True
run.font.name = 'Arial'

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("This Privacy Policy is effective as of " + today + ".")
run.italic = True
run.font.name = 'Arial'

# Save document
output_path = os.path.join(output_dir, "EzoraAI_Privacy_Policy.docx")
doc.save(output_path)
print("Privacy Policy document created successfully!")
print("  Location: " + output_path)
print("  Effective Date: " + today)
