#!/usr/bin/env python3
"""
Create EzoraAI Expert Agreement DOCX document using docx library
"""
import sys
sys.path.insert(0, '/sessions/keen-epic-bardeen/mnt/.claude/skills/docx')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set up page margins (1 inch)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "EzoraAI — Expert Agreement"
    header_para.runs[0].font.size = Pt(10)

    # Add footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    run = footer_para.add_run()
    run.font.size = Pt(10)

# Draft notice
draft = doc.add_paragraph()
draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
draft_run = draft.add_run("DRAFT — FOR LEGAL REVIEW")
draft_run.bold = True
draft_run.font.color.rgb = RGBColor(204, 0, 0)
draft_run.font.size = Pt(12)
draft.paragraph_format.space_after = Pt(18)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("EXPERT AGREEMENT & INDEPENDENT CONTRACTOR TERMS")
title_run.bold = True
title_run.font.size = Pt(16)
title.paragraph_format.space_after = Pt(12)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("EzoraAI Inc.")
subtitle_run.font.size = Pt(12)
subtitle.paragraph_format.space_after = Pt(24)

# Introduction
intro = doc.add_paragraph()
intro_run = intro.add_run('This Expert Agreement ("Agreement") is entered into between ')
intro_run.font.size = Pt(12)

# Section 1
heading1 = doc.add_heading("1. PARTIES", level=2)
heading1.paragraph_format.space_before = Pt(18)
heading1.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run('This Agreement is entered into between ')
p_run.font.size = Pt(12)
bold_run = p.add_run('EzoraAI Inc.')
bold_run.font.size = Pt(12)
bold_run.bold = True
final_run = p.add_run(', a Delaware corporation ("Platform" or "Company"), and the undersigned expert ("Expert").')
final_run.font.size = Pt(12)

doc.add_paragraph()

# Section 2
heading2 = doc.add_heading("2. INDEPENDENT CONTRACTOR STATUS", level=2)
heading2.paragraph_format.space_before = Pt(18)
heading2.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run("Expert is an independent contractor, NOT an employee, partner, joint venturer, or agent of the Company. Specifically:")
p_run.font.size = Pt(12)

bullets = [
    "Expert is solely responsible for all payroll taxes, Social Security, Medicare, unemployment insurance, and workers' compensation insurance",
    "Expert receives no employee benefits: no health insurance, retirement plans, paid time off, or other benefits",
    "Expert controls the methods, schedule, location, and manner of work delivery",
    "Expert sets their own rates (within Platform guidelines) and controls availability",
    "For US-based Experts: Company will issue Form 1099-NEC at year-end for tax reporting",
    "Expert is responsible for maintaining professional liability insurance, if desired"
]

for bullet in bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_paragraph()

# Section 3
heading3 = doc.add_heading("3. EXPERT OBLIGATIONS", level=2)
heading3.paragraph_format.space_before = Pt(18)
heading3.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run("Expert agrees to:")
p_run.font.size = Pt(12)

obligations = [
    ("Accurate Representation", "Truthfully represent qualifications, experience, certifications, and skills on profile. Misrepresentation may result in account suspension or termination."),
    ("Professional Conduct", "Maintain professional demeanor in all sessions, communications, and interactions. No harassment, discrimination, or abusive behavior."),
    ("Quality Standards", "Deliver high-quality expertise, meet client expectations, and provide value during all sessions."),
    ("Timely Responses", "Respond to booking requests and client inquiries within 48 hours."),
    ("Project Delivery", "Complete agreed projects or deliverables by the mutually agreed deadline. Extensions require prior written approval."),
    ("Compliance", "Comply with all Platform policies, including content guidelines, acceptable use policies, and code of conduct.")
]

for i, (title_text, desc) in enumerate(obligations, 1):
    p = doc.add_paragraph(style='List Number')
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Section 4
heading4 = doc.add_heading("4. PLATFORM FEES & TAKE RATE", level=2)
heading4.paragraph_format.space_before = Pt(18)
heading4.paragraph_format.space_after = Pt(9)

sections_4 = [
    ("4.1 Standard Take Rate", "Company retains 15–20% of session fees and project revenue as its platform fee. Expert receives the remaining 80–85%."),
    ("4.2 Founding Expert Rate", "Experts who reach 'Founding Expert' status (to be defined by Company) qualify for a reduced 10% take rate for 12 months from qualification date."),
    ("4.3 Rate Changes", "Company may adjust the take rate with 60 days' prior written notice. Expert may accept the new rate or terminate this Agreement."),
    ("4.4 Premium Features", "Company may offer optional premium features (featured listings, advanced analytics, certified badges) at additional cost. These are separate from session fees.")
]

for title_text, desc in sections_4:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 5
heading5 = doc.add_heading("5. PAYMENT TERMS", level=2)
heading5.paragraph_format.space_before = Pt(18)
heading5.paragraph_format.space_after = Pt(9)

sections_5 = [
    ("5.1 Payment Method", "All payments are processed via Stripe Connect. Expert must maintain an active Stripe account and valid payment method on file."),
    ("5.2 Payout Schedule", "Payouts occur weekly (Mondays) for sessions completed 7+ days prior. A 7-day hold is maintained on all transactions to address disputes and chargebacks."),
    ("5.3 Currency & International", "Payments are in USD. International Experts are responsible for currency conversion and international payment fees charged by their bank or Stripe."),
    ("5.4 Stripe Account Setup", "Expert is responsible for establishing and maintaining Stripe Connect account. Company is not liable for payment delays due to Expert's account issues, verification failures, or compliance violations.")
]

for title_text, desc in sections_5:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 6
heading6 = doc.add_heading("6. RATE SETTING & PRICING", level=2)
heading6.paragraph_format.space_before = Pt(18)
heading6.paragraph_format.space_after = Pt(9)

sections_6 = [
    ("6.1 Expert-Controlled Pricing", "Expert sets their own session rates within Company guidelines. At launch, the guideline range is $40–$100 per session."),
    ("6.2 Rate Adjustments", "Expert may adjust rates at any time. Changes take effect for new bookings within 24 hours of submission."),
    ("6.3 Platform Rate Range Changes", "Company may adjust the acceptable rate range ($40–$100) with 30 days' notice. Experts whose rates fall outside the new range will be notified and may adjust or request exception approval."),
    ("6.4 No Price Fixing", "Expert agrees not to collude with other Experts to set prices, undercut competitors, or engage in anti-competitive practices.")
]

for title_text, desc in sections_6:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 7
heading7 = doc.add_heading("7. INTELLECTUAL PROPERTY", level=2)
heading7.paragraph_format.space_before = Pt(18)
heading7.paragraph_format.space_after = Pt(9)

sections_7 = [
    ("7.1 Expert Content Ownership", "Expert retains all intellectual property rights to their own materials, tools, frameworks, templates, and content created prior to or outside the Platform."),
    ("7.2 Platform License", "Expert grants Company a non-exclusive, royalty-free license to display Expert's profile, qualifications, credentials, ratings, reviews, and work samples on the Platform."),
    ("7.3 AI Agents & Tools", "If Expert sells AI agents, custom tools, or software through the Platform, Expert retains ownership and grants Company a distribution license. Company takes its standard fee on sales."),
    ("7.4 Session Recordings", "Unless otherwise agreed in writing, Expert consents to recording sessions for quality assurance and record-keeping purposes. Recordings are confidential and not shared with third parties without consent.")
]

for title_text, desc in sections_7:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 8
heading8 = doc.add_heading("8. BLOCKCHAIN CREDENTIALS & ON-CHAIN VERIFICATION", level=2)
heading8.paragraph_format.space_before = Pt(18)
heading8.paragraph_format.space_after = Pt(9)

sections_8 = [
    ("8.1 Consent to On-Chain Recording", "Expert consents to Company recording verified credentials and work history on a blockchain ledger. This creates an immutable proof-of-work record of expertise."),
    ("8.2 What Goes On-Chain", "Records may include: (a) Expert profile verification date, (b) Session completion dates and client ratings, (c) Skill certifications and badges, (d) Total sessions completed, (e) Average client ratings."),
    ("8.3 Immutability", "Expert acknowledges that blockchain records cannot be deleted, modified, or revoked once recorded. Records are permanently available and publicly verifiable on the blockchain."),
    ("8.4 Privacy", "Blockchain records do not include sensitive client information (names, payment details, session content). Only Expert profile data and anonymized session metadata are recorded.")
]

for title_text, desc in sections_8:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 9
heading9 = doc.add_heading("9. NON-CIRCUMVENTION", level=2)
heading9.paragraph_format.space_before = Pt(18)
heading9.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run("Expert agrees not to solicit, book, or conduct business with any Platform client outside the Platform for 12 months after the Expert's last session with that client. This protects the Platform's investment in marketing and client acquisition.")
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Exception: ")
p_run.bold = True
p_run.font.size = Pt(12)
exception_run = p.add_run("This restriction does not apply to pre-existing client relationships that existed before the Expert joined the Platform, provided Expert discloses this relationship in writing at onboarding.")
exception_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Violation of this clause may result in immediate account suspension, reversal of pending payments, and Platform referral to legal enforcement.")
p_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(18)

doc.add_paragraph()

# Section 10
heading10 = doc.add_heading("10. CONFIDENTIALITY", level=2)
heading10.paragraph_format.space_before = Pt(18)
heading10.paragraph_format.space_after = Pt(9)

sections_10 = [
    ("10.1 Client Information", "Expert will not disclose client names, contact details, session content, or project information to third parties without prior written consent from the client."),
    ("10.2 Platform Data", "Expert will not share Platform proprietary information, algorithms, fee structures, client lists, or internal data with competitors or the public."),
    ("10.3 Duration", "These confidentiality obligations survive termination of this Agreement indefinitely.")
]

for title_text, desc in sections_10:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 11
heading11 = doc.add_heading("11. TERMINATION", level=2)
heading11.paragraph_format.space_before = Pt(18)
heading11.paragraph_format.space_after = Pt(9)

sections_11 = [
    ("11.1 Termination Without Cause", "Either party may terminate this Agreement by providing 14 days' written notice. Expert will stop accepting new bookings upon notice and fulfill existing commitments."),
    ("11.2 Immediate Termination", "Company may immediately suspend or terminate Expert's account for: (a) policy violations, (b) harmful conduct, (c) fraudulent activity, (d) repeated complaints or poor ratings, (e) breach of this Agreement."),
    ("11.3 Outstanding Payments", "Company will process all earned payments within 30 days of termination, minus any refunds due to disputes or chargebacks."),
    ("11.4 Surviving Clauses", "Sections 7 (IP), 8 (Blockchain), 9 (Non-Circumvention), 10 (Confidentiality), 12 (Limitation of Liability), and 13 (Dispute Resolution) survive termination.")
]

for title_text, desc in sections_11:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 12
heading12 = doc.add_heading("12. LIMITATION OF LIABILITY", level=2)
heading12.paragraph_format.space_before = Pt(18)
heading12.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run("TO THE MAXIMUM EXTENT PERMITTED BY LAW, NEITHER PARTY SHALL BE LIABLE FOR INDIRECT, INCIDENTAL, CONSEQUENTIAL, SPECIAL, OR PUNITIVE DAMAGES, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGES. COMPANY'S TOTAL LIABILITY UNDER THIS AGREEMENT SHALL NOT EXCEED THE FEES PAID BY THE EXPERT IN THE 12 MONTHS PRECEDING THE CLAIM.")
p_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(18)

doc.add_paragraph()

# Section 13
heading13 = doc.add_heading("13. DISPUTE RESOLUTION", level=2)
heading13.paragraph_format.space_before = Pt(18)
heading13.paragraph_format.space_after = Pt(9)

sections_13 = [
    ("13.1 Mediation", "Before litigation or arbitration, the parties agree to attempt good-faith mediation. Either party may initiate mediation by submitting a written request to the other party."),
    ("13.2 Arbitration", "If mediation fails, any dispute shall be resolved by binding arbitration administered by JAMS (Judicial Arbitration and Mediation Services) or similar, in Delaware, under JAMS Rules. Arbitration is confidential and final."),
    ("13.3 No Class Action", "Expert waives the right to bring class action claims. All disputes are resolved individually.")
]

for title_text, desc in sections_13:
    p = doc.add_paragraph()
    bold_run = p.add_run(f"{title_text}: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(9)

doc.add_paragraph()

# Section 14
heading14 = doc.add_heading("14. GOVERNING LAW", level=2)
heading14.paragraph_format.space_before = Pt(18)
heading14.paragraph_format.space_after = Pt(9)

p = doc.add_paragraph()
p_run = p.add_run("This Agreement is governed by and construed in accordance with the laws of the State of Delaware, without regard to conflicts of law principles.")
p_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(24)

# Page break
doc.add_page_break()

# Signature section
heading_sig = doc.add_heading("SIGNATURE BLOCK", level=2)
heading_sig.paragraph_format.space_before = Pt(18)
heading_sig.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph()
p_run = p.add_run("By signing below, both parties agree to the terms of this Expert Agreement.")
p_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(18)

# Company signature section
p = doc.add_paragraph()
bold_run = p.add_run("FOR EZORAIINC.")
bold_run.bold = True
bold_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Authorized Representative (Print Name):")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Signature:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Title:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Date:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(18)

# Expert signature section
p = doc.add_paragraph()
bold_run = p.add_run("FOR EXPERT")
bold_run.bold = True
bold_run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Expert Name (Print):")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Signature:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Date:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run("Email:")
p_run.font.size = Pt(12)
p = doc.add_paragraph()
p_run = p.add_run("_" * 70)
p_run.font.size = Pt(12)

# Save document
output_path = "/sessions/keen-epic-bardeen/mnt/EzoraAI-Ops/EzoraAI/Legal/EzoraAI_Expert_Agreement.docx"
doc.save(output_path)
print(f"✓ Expert Agreement created successfully at {output_path}")
